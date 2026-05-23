"""Generate the Regime Compass methodology report as a Word document."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "Regime_Compass_Methodology.docx"


def _set_run(run, *, bold=False, italic=False, size=11, color=None, mono=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    if mono:
        run.font.name = "Menlo"


def _add_h1(doc, text):
    p = doc.add_heading("", level=1)
    r = p.add_run(text)
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x16, 0x16, 0x32)


def _add_h2(doc, text):
    p = doc.add_heading("", level=2)
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x66)


def _add_para(doc, text, *, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, italic=italic, size=size)


def _add_rich_para(doc, parts, *, size=11):
    """Add a paragraph with mixed formatting. `parts` is a list of (text, {**fmt})."""
    p = doc.add_paragraph()
    for part in parts:
        text, fmt = part if isinstance(part, tuple) else (part, {})
        r = p.add_run(text)
        _set_run(r, size=size, **fmt)


def _add_bullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, str):
            r = p.add_run(it)
            _set_run(r, size=size)
        else:
            for part in it:
                text, fmt = part if isinstance(part, tuple) else (part, {})
                r = p.add_run(text)
                _set_run(r, size=size, **fmt)


def _add_table(doc, header, rows, *, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        _set_run(r, bold=True, size=10)
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            _set_run(r, size=10)
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)


def build():
    doc = Document()
    # Slightly tighter margins for 4-page fit
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Title
    title = doc.add_paragraph()
    r = title.add_run("Regime Compass")
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x16, 0x16, 0x32)
    r.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph()
    r = subtitle.add_run("Statistical Methodology & Product Overview")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x6b, 0x74, 0x80)
    r.italic = True

    _add_para(
        doc,
        "A free, daily-updated market regime classifier across six global markets. "
        "Built by Aditya Sahasrabuddhe under the iQuant Labs umbrella. "
        "Live at regimecompass.com.",
        italic=True,
        size=10,
    )
    doc.add_paragraph()  # spacer

    # Section 1
    _add_h2(doc, "1.  What Regime Compass is, in one paragraph")
    _add_para(
        doc,
        "Regime Compass is a free web tool that classifies the current state of six global markets — "
        "Nifty 50, S&P 500, KOSPI, Shanghai Composite, Bitcoin, and Ethereum — into one of three "
        "regimes: bull, neutral, or bear. It does this through three independent statistical methods "
        "running side by side: a 3-state Gaussian Hidden Markov Model, simple moving-average crossovers "
        "(50 / 100 / 200-day), and exponential moving-average crossovers (same periods). The site is "
        "updated automatically every trading day, exposes 16 years of historical regime classifications, "
        "runs honest walk-forward backtests of the most common timing strategies, and offers email "
        "alerts when a market changes regime."
    )
    _add_rich_para(doc, [
        "It is explicitly ",
        ("not", {"italic": True}),
        " a trade-signal generator, prediction system, or robo-advisor. It is a ",
        ("descriptive indicator", {"bold": True}),
        " — it tells you what kind of market you're trading right now, not what to do about it. "
        "The interpretation is left to the human."
    ])

    # Section 2
    _add_h2(doc, "2.  What problem it tries to solve")
    _add_para(
        doc,
        "Most discretionary traders intuit market regimes — \"we're in a risk-off period,\" "
        "\"vol is breaking out,\" \"this is a low-vol grind\" — but their intuition is unstructured "
        "and often lags reality by weeks. Equally, most quantitative regime tools that do exist are "
        "either:"
    )
    _add_bullets(doc, [
        "Black-box (Bloomberg / Refinitiv terminals at $20K+/year, no exposed methodology)",
        "One-market-only (e.g., US-centric CNN Fear & Greed)",
        "Strategy-disguised-as-indicator (sell you a \"regime\" label that's really a buy/sell signal in disguise)",
        "Free but stale (built on outdated public datasets, no daily refresh)",
    ])
    _add_rich_para(doc, [
        "Regime Compass aims to be the ",
        ("transparent, multi-market, daily-updated middle ground", {"bold": True}),
        ": methods fully documented, formulas in the open, no signup gate, six markets including "
        "crypto, and no claims beyond what the math actually supports."
    ])

    # Section 3
    _add_h2(doc, "3.  The statistical logic")

    _add_h2(doc, "3.1  Hidden Markov Model (the headline method)")
    _add_para(
        doc,
        "The site fits a 3-state Gaussian Hidden Markov Model independently per market using the "
        "hmmlearn library, trained on approximately 16 years of daily data (2010 — present, except "
        "Bitcoin and Ethereum, which start in 2014 and 2017 respectively)."
    )
    _add_para(
        doc,
        "The model assumes that on each trading day, the market is in one of three unobserved states, "
        "and the observed features for that day are drawn from a multivariate Gaussian distribution "
        "specific to that state. Transitions between states follow a Markov chain. Formally, let S_t "
        "denote the hidden regime at time t; then P(S_t | S_{t-1}) is the transition matrix "
        "(estimated via Baum-Welch / EM), and X_t | S_t = s is drawn from a Gaussian with mean μ_s "
        "and covariance Σ_s."
    )
    _add_para(doc, "Feature vector (4 dimensions where available, 3 otherwise):", size=10)
    _add_bullets(doc, [
        "Daily log return: ln(P_t / P_{t-1})",
        "10-day rolling realized volatility: standard deviation of the previous 10 log returns",
        "FX log-change vs USD (USD/INR for Nifty, DXY for SPX, etc.)",
        "Implied volatility index level — India VIX for Nifty, CBOE VIX for SPX. Not available for KOSPI, Shanghai, BTC, or ETH; those models use 3 features.",
    ])
    _add_rich_para(doc, [
        "All features are ",
        ("z-scored before fitting", {"bold": True}),
        " so no single feature dominates by magnitude."
    ])

    _add_h2(doc, "3.2  Stable state labeling (the design choice that matters)")
    _add_para(
        doc,
        "The HMM doesn't know which of its three states is \"bull\" or \"bear\" — labels are "
        "arbitrary. After fitting, we sort the three states by a composite score and assign labels "
        "deterministically:"
    )
    p = doc.add_paragraph()
    r = p.add_run("score(s) = − z(mean_vol_s) + 0.25 × z(mean_return_s)")
    _set_run(r, mono=True, size=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_rich_para(doc, [
        "where z(·) is z-scoring across the three states. The lowest score becomes ",
        ("bear", {"bold": True}),
        ", the highest becomes ",
        ("bull", {"bold": True}),
        ", the middle is neutral."
    ])

    _add_rich_para(doc, [
        ("Why vol-dominant?", {"bold": True}),
        " An earlier version used a balanced rule (score = z(return) − z(vol)). It failed the "
        "eye-test catastrophically: COVID March 2020 was classified as \"neutral\" with 100% "
        "probability, while sideways grinds in 2017 got labeled \"bear.\" This is a known failure "
        "mode for return-balanced labeling on long-bull equity markets: with 15+ years of net upward "
        "drift, mean returns barely separate states. Volatility is the actionable axis. We give "
        "returns a small weight (0.25) as a tiebreaker."
    ])

    _add_h2(doc, "3.3  Filtered, not smoothed, probabilities")
    _add_para(
        doc,
        "A subtle but critical distinction. The model produces a probability distribution over states "
        "for every historical day. Two ways to compute these:"
    )
    _add_bullets(doc, [
        [("Smoothed: ", {"bold": True}),
         "P(S_t | X_1, …, X_T). Uses both past and future observations. Produces beautiful historical "
         "charts but cheats; can't be computed in real time."],
        [("Filtered: ", {"bold": True}),
         "P(S_t | X_1, …, X_t). Uses only data available up to day t. Lags slightly but is causal."],
    ])
    _add_rich_para(doc, [
        ("Regime Compass always displays filtered probabilities", {"bold": True}),
        " on the live dashboard. Smoothed probabilities exist only in the offline analysis notebooks. "
        "This was verified by independent re-implementation against scipy.special.logsumexp (matches to "
        "10⁻¹³ precision)."
    ])

    _add_h2(doc, "3.4  Moving-average regimes (the deterministic baselines)")
    _add_para(
        doc,
        "For each market, we also compute simple price-vs-moving-average regime classifications at three "
        "timeframes: 50-day, 100-day, and 200-day, in both Simple Moving Average (SMA) and Exponential "
        "Moving Average (EMA) flavors."
    )
    _add_bullets(doc, [
        [("SMA rule: ", {"bold": True}), "price > MA → bull; price < MA → bear."],
        [("EMA rule: ", {"bold": True}),
         "identical, but uses EMA_t = α × P_t + (1 − α) × EMA_{t−1} with α = 2 / (N+1)."],
    ])
    _add_para(
        doc,
        "These are intentionally simple. They serve three purposes: (a) sanity-check the HMM (if all "
        "three disagree, something's off); (b) act as low-noise trend filters for traders who don't "
        "trust machine-learning methods; (c) feed into the composite risk score."
    )

    _add_h2(doc, "3.5  The per-market risk score")
    _add_para(
        doc,
        "For each market, we compute a single scalar in [0, 100] blending the three signals:"
    )
    _add_bullets(doc, [
        "HMM score: (P_bull − P_bear) × 100",
        "200-SMA score: +100 if bull, −100 if bear",
        "200-EMA score: same",
    ])
    _add_para(
        doc,
        "The market's score is the arithmetic mean of the three sub-scores, then linearly rescaled to "
        "a 0–100 gauge. Above 70 = risk-on (all three models agree the market is bullish). Below 30 = "
        "risk-off (all three agree bearish). Between 30 and 70 = mixed (signals disagree, often around "
        "turning points)."
    )
    _add_rich_para(doc, [
        "Notably, the site ",
        ("does not produce a global composite across all six markets", {"bold": True}),
        ". Aggregating six index scores into a single \"world risk score\" with this small a sample isn't "
        "statistically meaningful, and we say so explicitly on the page. Each market is shown independently."
    ])

    _add_h2(doc, "3.6  Backtests (walk-forward, 2-day confirmation)")
    _add_para(
        doc,
        "For the moving-average models, we run a walk-forward backtest on each market with a 2-day "
        "confirmation rule:"
    )
    _add_bullets(doc, [
        "Long the index from close[t−1] to close[t] when the previous two consecutive days showed bull regime",
        "Hold cash earning the country's approximate risk-free rate (6.5% India, 2% US, 2.5% Korea/China, 2% USD-denominated crypto) when the previous two days showed bear",
    ])
    _add_para(
        doc,
        "The 2-day rule reduces whipsaws from single-day regime flips. Position decisions use only "
        "past data (truly walk-forward), no transaction costs, no slippage, no borrow costs. Returns "
        "are pre-cost gross."
    )
    _add_rich_para(doc, [
        ("The HMM backtest was deliberately removed.", {"bold": True}),
        " Simple binary HMM timing underperforms buy-and-hold by 40–50% on absolute return across "
        "every market we tested — a structural truth of timing strategies in long-bull markets, not "
        "a model defect. Displaying a flawed strategy in our UI would mislead users."
    ])
    _add_rich_para(doc, [
        ("What the backtests actually show:", {"bold": True}),
        " the regime-overlay strategies tend to give up 30–60% of absolute return vs. buy-and-hold, "
        "but cut maximum drawdown by 40–60%. Sharpe ratios are roughly equal. The value proposition "
        "isn't \"more money\"; it's \"the same risk-adjusted return with materially less stress.\""
    ])

    # Section 4
    _add_h2(doc, "4.  Forward conditional returns — the most useful output")
    _add_para(
        doc,
        "For each historical day, we observe what the market actually did over the next 5 / 30 / 60 / "
        "90 trading days, then group by the regime the model assigned that day. This is descriptive, "
        "not predictive — it answers the question, \"historically, when the model said X regime, what "
        "tended to happen next?\""
    )
    _add_para(doc, "The surprising finding, which holds across markets:")
    _add_table(
        doc,
        ["Regime today (Nifty, 16yr)", "Avg forward 90-day return", "Win rate"],
        [
            ["Bear", "+8.5%", "78.7%"],
            ["Neutral", "+3.7%", "64.5%"],
            ["Bull", "+2.6%", "68.3%"],
        ],
        col_widths=[6.5, 5.5, 3.5],
    )
    _add_para(
        doc,
        "Bear regimes historically had the highest forward returns, not the lowest. The same pattern "
        "holds on S&P 500 (+7.8% / 81%), KOSPI (+11.7% / 85%). The interpretation: bear regimes are "
        "triggered by vol spikes which typically arrive after the drop, so the model flags panic just "
        "as recovery is statistically about to begin."
    )
    _add_rich_para(doc, [
        ("The naive \"bear = sell\" interpretation is exactly backward.", {"bold": True}),
        " The honest use of the indicator is closer to \"bear = context for contrarian sizing.\" This "
        "is the unique insight Regime Compass surfaces that most timing tools do not."
    ])

    # Section 5
    _add_h2(doc, "5.  Data infrastructure and update cadence")
    _add_bullets(doc, [
        [("Source: ", {"bold": True}), "Yahoo Finance via the yfinance library — free, daily-updated, accepted for non-commercial educational use"],
        [("Markets: ", {"bold": True}), "^NSEI (Nifty), ^GSPC (S&P 500), ^KS11 (KOSPI), 000001.SS (Shanghai), BTC-USD, ETH-USD"],
        [("Daily refresh: ", {"bold": True}), "APScheduler runs in-process at 11:00 UTC Mon–Fri (just after NSE close), fetching latest prices, recomputing filtered probabilities, dispatching email alerts"],
        [("Weekly retrain: ", {"bold": True}), "Sunday 03:30 UTC, full retrain of all 6 HMMs on the now-week-bigger dataset"],
        [("Resilience: ", {"bold": True}), "if a fetch fails, yesterday's data is preserved (atomic writes); a \"data delayed\" banner appears if data is more than 72 hours stale"],
        [("Bootstrap on restart: ", {"bold": True}), "when the container restarts, the site automatically refetches all data and retrains — guaranteeing freshness without persistent storage costs"],
    ])
    _add_para(
        doc,
        "A full audit harness (in the codebase) runs over 100 statistical assertions on every refit "
        "to verify features compute correctly, the scaler is invertible, HMM converges, label "
        "permutation is stable, filtered probabilities sum to 1, the forward algorithm matches an "
        "independent SciPy implementation, and the database matches in-memory state."
    )

    # Section 6
    _add_h2(doc, "6.  What you actually see when you visit")
    _add_bullets(doc, [
        [("Home (/): ", {"bold": True}), "snapshot grids for SMA, EMA, and HMM across all six markets — one cell per (market × timeframe) showing today's regime"],
        [("Risk Score (/composite): ", {"bold": True}), "six per-market gauges (0–100) blending the three model families"],
        [("Detail pages (/hmm, /ma, /ema): ", {"bold": True}), "per-market deep dives with charts, state means, regime histories, recent runs"],
        [("Backtests (/ma/backtest, /ema/backtest): ", {"bold": True}), "walk-forward backtest with strategy vs. buy-and-hold metrics, equity curve, trade table, drawdown reduction stat"],
        [("Alerts (/subscribe): ", {"bold": True}), "free email alerts when a regime flips on any subscribed market or when the per-market risk score crosses 30 or 70"],
        [("About / Methodology / Disclaimer / Privacy: ", {"bold": True}), "full documentation of every design choice; nothing hidden"],
    ])

    # Section 7
    _add_h2(doc, "7.  What it deliberately does not do")
    _add_bullets(doc, [
        "Does not predict prices",
        "Does not recommend trades",
        "Does not generate buy / sell signals",
        "Does not aggregate markets into a single global \"world fear\" index",
        "Does not paywall any feature, even at scale",
        "Does not collect personal data beyond what's needed to send a verified email alert",
    ])

    # Section 8
    _add_h2(doc, "8.  Honest limitations")
    _add_bullets(doc, [
        "HMMs lag regime changes by 1–3 trading days on average — by the time the model confirms a regime, the move has begun",
        "Free data (yfinance) is occasionally noisy; an audit found 4 days of suspect USD/INR ticks in early 2012 and 2 days of suspect USD/CNY in mid-2011, both with bounded local impact",
        "The model is blind to information not in its features — fundamentals, flows, sentiment, news, macro releases are invisible",
        "Three states is a deliberate simplification; real markets have more nuance",
        "All conclusions assume statistical stationarity over the training window, which is itself an assumption that can fail at structural regime shifts",
    ])

    # Section 9
    _add_h2(doc, "9.  Roadmap")
    _add_para(
        doc,
        "Planned future additions: a sentiment-indicators page that surfaces existing third-party "
        "sentiment data (CNN Fear & Greed, Crypto Fear & Greed, FII/DII flows from NSE) rather than "
        "building a proprietary composite; a sector-level regime heatmap for Nifty sectoral indices; "
        "a portfolio analyzer that takes a holdings CSV and reports the regime breakdown of the user's "
        "actual positions."
    )

    # Closing line
    doc.add_paragraph()
    _add_para(
        doc,
        "This document and the underlying tool are educational and informational only. They are not "
        "investment advice in any jurisdiction. The author is not a registered investment advisor "
        "anywhere in the world. Trading and investing in financial markets carries substantial risk "
        "of loss including total loss of capital.",
        italic=True,
        size=9,
    )
    _add_para(
        doc,
        "— Aditya Sahasrabuddhe · linkedin.com/in/aditya-s1/",
        italic=True,
        size=9,
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
